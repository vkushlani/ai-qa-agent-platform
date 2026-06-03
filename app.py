import time
from pathlib import Path

import streamlit as st
import pandas as pd
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

from app_router import classify_question
from app_agents import coordinator_agent
from app_memory_manager import (
    save_memory,
    load_memories,
    save_memory_to_vector_db,
    retrieve_memory,
)

# =====================================================
# ENVIRONMENT
# =====================================================

env_path = Path(__file__).with_name(".env")
load_dotenv(dotenv_path=env_path)

# =====================================================
# PAGE CONFIG
# =====================================================

st.set_page_config(
    page_title="AI-Powered QA Agent Platform",
    layout="wide"
)

st.title("AI-Powered QA Agent Platform")

st.caption(
    "AI assistant for test design, defect analysis, traceability, regression risk assessment, document analysis, and QA planning."
)

st.markdown("""
This app demonstrates an AI assistant for software testing.

You can use it to:

- Generate test cases from requirements
- Analyze defects and root causes
- Create requirement traceability ideas
- Identify regression risks
- Run mock automation checks

You may ask questions directly, or optionally upload documents such as:
release notes, requirements, defect reports, test plans, or user stories.
""")

# =====================================================
# LLM
# =====================================================

llm = ChatOpenAI(
    model="gpt-4.1-mini",
    temperature=0.2
)

# =====================================================
# SESSION STATE
# =====================================================

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "uploaded_file_names" not in st.session_state:
    st.session_state.uploaded_file_names = []

if "uploaded_file_contents" not in st.session_state:
    st.session_state.uploaded_file_contents = {}

past_memories = load_memories()

DEBUG_MODE = False

# =====================================================
# FILE EXTRACTION
# =====================================================

def extract_file_content(uploaded_file):

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    elif file_name.endswith(".pdf"):
        pdf_reader = PdfReader(uploaded_file)
        text = ""

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        return text

    elif file_name.endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text

    elif file_name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
        return df.to_string(index=False)

    elif file_name.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
        return df.to_string(index=False)

    return ""

# =====================================================
# OPTIONAL DOCUMENT UPLOAD
# =====================================================

uploaded_files = st.file_uploader(
    "Optional: Upload QA Documents",
    type=["txt", "pdf", "docx", "csv", "xlsx"],
    accept_multiple_files=True,
    help="Upload release notes, requirements, defect reports, test plans, user stories, Excel, CSV, Word, or PDF files."
)

if uploaded_files:

    st.session_state.uploaded_file_names = []
    st.session_state.uploaded_file_contents = {}

    for uploaded_file in uploaded_files:

        file_name = uploaded_file.name
        file_content = extract_file_content(uploaded_file)

        st.session_state.uploaded_file_names.append(file_name)
        st.session_state.uploaded_file_contents[file_name] = file_content

    uploaded_names = ", ".join(
        st.session_state.uploaded_file_names
    )

    st.success(
        f"✅ {len(st.session_state.uploaded_file_names)} document(s) uploaded: {uploaded_names}"
    )

# =====================================================
# DISPLAY CHAT HISTORY
# =====================================================

st.subheader("Conversation")

for message in st.session_state.chat_history:

    if isinstance(message, dict):

        with st.chat_message(message.get("role", "assistant")):

            st.write(message.get("content", ""))

            if (
                message.get("role") == "assistant"
                and "response_time" in message
            ):
                st.caption(
                    f"⏱️ Response generated in {message['response_time']:.2f} seconds"
                )

# =====================================================
# CHAT INPUT
# =====================================================

question = st.chat_input("Ask the AI Testing Agent...")

# =====================================================
# MAIN PROCESS
# =====================================================

if question:

    start_time = time.time()

    # Show current user message immediately in this run
    with st.chat_message("user"):
        st.write(question)

    with st.chat_message("assistant"):

        with st.spinner("Thinking..."):

            query_type = classify_question(question)

            document_required_workflows = [
                "comparison",
                "summary",
                "document_count",
                "document_names",
            ]

            if (
                query_type in document_required_workflows
                and
                not st.session_state.uploaded_file_names
            ):

                answer = (
                    "This request requires uploaded documents. "
                    "Please upload release notes, requirements, defect reports, "
                    "test plans, or similar QA documents."
                )

            else:

                # =================================================
                # INITIALIZE VARIABLES
                # =================================================

                all_documents = []
                uploaded_document_names = []
                combined_content = ""
                unique_documents = []
                retrieved_context = ""
                memory_context = ""

                history = "\n".join(
                    [
                        f"{msg.get('role', '')}: {msg.get('content', '')}"
                        for msg in st.session_state.chat_history
                        if isinstance(msg, dict)
                    ]
                )

                # =================================================
                # READ DOCUMENTS FROM SESSION STATE
                # =================================================

                if st.session_state.uploaded_file_contents:

                    for filename, content in st.session_state.uploaded_file_contents.items():

                        uploaded_document_names.append(filename)

                        combined_content += f"""

DOCUMENT NAME:
{filename}

DOCUMENT CONTENT:
{content}

"""

                        all_documents.append(
                            Document(
                                page_content=content,
                                metadata={"source": filename}
                            )
                        )

                    unique_documents = list(
                        set(uploaded_document_names)
                    )

                # =================================================
                # MEMORY RETRIEVAL
                # Disable memory for document metadata workflows
                # =================================================

                if query_type in [
                    "document_count",
                    "document_names",
                    "comparison",
                    "summary",
                ]:
                    memory_results = []
                else:
                    memory_results = retrieve_memory(question)

                for memory in memory_results:

                    memory_context += f"""

{memory.page_content}

"""

                # =================================================
                # RAG ONLY IF DOCUMENTS EXIST
                # =================================================

                if all_documents:

                    splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=80
                    )

                    chunks = splitter.split_documents(all_documents)

                    embedding_model = OpenAIEmbeddings()

                    vectorstore = Chroma.from_documents(
                        documents=chunks,
                        embedding=embedding_model,
                        persist_directory="vector_db"
                    )

                    results = vectorstore.similarity_search(
                        question,
                        k=6
                    )

                    for r in results:

                        source = r.metadata.get("source", "Unknown")

                        retrieved_context += f"""

DOCUMENT:
{source}

CONTENT:
{r.page_content}

"""

                # =================================================
                # MASTER CONTEXT
                # =================================================

                master_context = f"""

Uploaded Documents:
{unique_documents}

Conversation History:
{history}

Historical Memory:
{memory_context}

Retrieved Context:
{retrieved_context}

Document Content:
{combined_content}

"""

                # =================================================
                # DOCUMENT COUNT
                # =================================================

                if query_type == "document_count":

                    answer = f"""
Total uploaded documents: {len(unique_documents)}

Document names:

{chr(10).join(unique_documents)}
"""

                # =================================================
                # DOCUMENT NAMES
                # =================================================

                elif query_type == "document_names":

                    answer = f"""
Uploaded document name(s):

{chr(10).join(unique_documents)}
"""

                # =================================================
                # SUMMARY
                # =================================================

                elif query_type == "summary":

                    prompt = f"""
You are a document summarization expert.

Summarize all uploaded documents clearly.

Documents:
{unique_documents}

Document Content:
{combined_content}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

                # =================================================
                # COMPARISON
                # =================================================

                elif query_type == "comparison":

                    prompt = f"""
You are a document comparison expert.

Compare all uploaded documents.

Include:
1. Similarities
2. Differences
3. Key findings
4. Missing or unique topics

Documents:
{unique_documents}

Document Content:
{combined_content}

Question:
{question}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

                # =================================================
                # MULTI-AGENT WORKFLOWS
                # =================================================

                elif query_type in [
                    "test_case",
                    "defect_analysis",
                    "traceability",
                    "regression_risk",
                    "coverage_pipeline",
                    "automation",
                    "website_testing",
                    "planning",
                    "release_readiness",
                ]:

                    answer = coordinator_agent(
                        query_type=query_type,
                        context=master_context,
                        question=question
                    )

                # =================================================
                # DEFAULT GENERAL QA / RAG WORKFLOW
                # =================================================

                else:

                    if retrieved_context:

                        prompt = f"""
You are a Senior QA AI Assistant.

Use uploaded documents if relevant.

Conversation History:
{history}

Historical Memory:
{memory_context}

Retrieved Context:
{retrieved_context}

Question:
{question}
"""

                    else:

                        prompt = f"""
You are a friendly AI QA Assistant.

If the user is chatting casually, respond naturally.

If the user asks QA or software testing questions, answer using your QA knowledge.

Conversation History:
{history}

Historical Memory:
{memory_context}

Question:
{question}
"""

                    response = llm.invoke(prompt)
                    answer = response.content

            # =================================================
            # RESPONSE TIME
            # =================================================

            end_time = time.time()
            response_time = end_time - start_time

            # =================================================
            # DISPLAY ANSWER
            # =================================================

            st.write(answer)

            st.caption(
                f"⏱️ Response generated in {response_time:.2f} seconds"
            )

            # =================================================
            # SAVE CHAT HISTORY
            # =================================================

            st.session_state.chat_history.append(
                {
                    "role": "user",
                    "content": question,
                }
            )

            st.session_state.chat_history.append(
                {
                    "role": "assistant",
                    "content": answer,
                    "response_time": response_time,
                }
            )

            # =================================================
            # SAVE MEMORY
            # =================================================

            save_memory(question, answer)

            save_memory_to_vector_db(question, answer)

            # =================================================
            # DEBUG SECTIONS
            # =================================================

            if DEBUG_MODE:

                with st.expander("Retrieved Document Context"):
                    st.text(retrieved_context)

                with st.expander("Retrieved Memory Context"):
                    st.text(memory_context)

                with st.expander("Persistent Memory Store"):
                    st.write(past_memories)

# =====================================================
# FOOTER
# =====================================================

st.markdown("---")

st.caption(
    "Built with Streamlit, LangChain, OpenAI, ChromaDB, and multi-agent QA architecture."
)